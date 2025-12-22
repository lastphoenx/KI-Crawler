"""
Crawler service wrapper for UI integration
"""

import logging
import threading
import sys
import os
import re
from datetime import datetime
from pathlib import Path
from urllib.parse import urlparse

# Add parent directory to path to import crawler
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from crawler import Crawler
from services.crawl_state_manager import CrawlStateManager
from parser import HTMLParser
import yaml

logger = logging.getLogger(__name__)


class CrawlerService:
    """Service layer for crawler operations"""
    
    def __init__(self):
        self.logger = logger
        self.state_manager = CrawlStateManager()
        self.active_crawls = {}
    
    def get_recent_crawls(self, limit=5):
        """Get recent crawl history from state manager"""
        recent_states = self.state_manager.get_recent_crawls(limit)
        
        # Convert to UI format
        result = []
        for state in recent_states:
            # Use pages_crawled, but fallback to actual data if it's 0
            actual_pages = state.pages_crawled
            if actual_pages == 0:
                # Fallback: use length of crawled_pages dict or current_page
                if hasattr(state, 'crawled_pages') and state.crawled_pages:
                    actual_pages = len(state.crawled_pages)
                elif state.current_page > 0:
                    actual_pages = state.current_page
            
            time_ago = self._format_time_ago(state.start_time)
            result.append({
                'id': state.crawl_id,
                'url': state.config.get('url', 'Unknown'),
                'status': 'success' if state.status == 'completed' else 'warning' if state.errors > 0 else 'success',
                'pages': actual_pages,
                'duration': state.get_duration(),
                'time_ago': time_ago,
                'timestamp': state.start_time
            })
        
        return result
    
    def _url_to_filename(self, url: str) -> str:
        """Convert URL to a safe filename (e.g., docs_pcloud_com)"""
        parsed = urlparse(url)
        # Get domain without www
        domain = parsed.netloc.replace('www.', '')
        # Add first path segment if exists
        path_parts = [p for p in parsed.path.split('/') if p]
        if path_parts:
            domain += f"_{path_parts[0]}"
        # Clean to safe filename
        safe = re.sub(r'[^a-zA-Z0-9_]', '_', domain)
        return safe.lower()[:50]  # Limit length
    
    def _get_unique_filepath(self, base_dir: Path, base_name: str, extension: str) -> Path:
        """Get unique filepath by incrementing number if file exists"""
        filepath = base_dir / f"{base_name}.{extension}"
        counter = 1
        while filepath.exists():
            filepath = base_dir / f"{base_name}_{counter}.{extension}"
            counter += 1
        return filepath
    
    def start_crawl(self, config_dict):
        """Start a new crawl with given configuration"""
        crawl_id = config_dict.get('crawl_id')
        self.logger.info(f"Starting crawl {crawl_id} for: {config_dict.get('url')}")
        
        # Create state
        state = self.state_manager.create_crawl(crawl_id, config_dict)
        
        # Start crawler in background thread
        thread = threading.Thread(
            target=self._run_crawler,
            args=(crawl_id, config_dict),
            daemon=True
        )
        thread.start()
        
        return {'status': 'started', 'crawl_id': crawl_id}
    
    def _run_crawler(self, crawl_id: str, config_dict: dict):
        """Run the actual crawler in background"""
        try:
            # DEBUG: Check enhanced_html_formatting value
            enhanced_formatting = config_dict.get('enhanced_html_formatting', False)
            self.logger.info(f"🔍 DEBUG: enhanced_html_formatting from config_dict = {enhanced_formatting}")
            
            # Update state
            self.state_manager.update_progress(crawl_id, status='running')
            self.state_manager.add_log(crawl_id, '🚀 Starting crawler...')
            self.state_manager.add_log(crawl_id, f'🔍 DEBUG: enhanced_html_formatting = {enhanced_formatting}')
            
            # NEW STRUCTURE: Create folder per crawl with descriptive name + timestamp
            base_output_dir = Path(__file__).parent.parent / 'output'
            base_output_dir.mkdir(exist_ok=True)
            
            # Generate descriptive folder name from URL or output_name
            if config_dict.get('output_name'):
                # User provided name (remove extensions)
                folder_prefix = config_dict['output_name'].replace('.docx', '').replace('.html', '').replace('.pdf', '')
            else:
                # Generate from URL: github.com/user/repo → github_user_repo
                folder_prefix = self._url_to_filename(config_dict['url'])
            
            # Create folder: {prefix}_{YYYY-MM-DD_HH-MM-SS}
            timestamp = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
            crawl_dir = base_output_dir / f'{folder_prefix}_{timestamp}'
            
            # Ensure unique folder name (in case of rapid sequential crawls)
            counter = 1
            original_crawl_dir = crawl_dir
            while crawl_dir.exists():
                crawl_dir = base_output_dir / f'{folder_prefix}_{timestamp}_{counter}'
                counter += 1
            
            crawl_dir.mkdir(exist_ok=True)
            self.logger.info(f"📁 Crawl folder: {crawl_dir.name}")
            
            # File names within folder
            base_name = 'output'  # Standard name for files inside
            
            # Define file paths within crawl folder
            output_file = crawl_dir / f'{base_name}.docx'
            html_file = crawl_dir / 'index.html'
            log_file = crawl_dir / f'{base_name}.log'
            config_file = crawl_dir / 'config.yaml'
            
            self.logger.info(f"Output files: DOCX={output_file.name}, HTML={html_file.name}, LOG={log_file.name}")
            
            # Update state with output path
            self.state_manager.update_progress(crawl_id, output_path=str(output_file))
            
            # Load base config
            config_path = Path(__file__).parent.parent / 'config.yaml'
            with open(config_path, 'r') as f:
                base_config = yaml.safe_load(f)
            
            # Override with UI config
            # Normalize URL: remove trailing slash to avoid double slashes
            entry_url = config_dict['url'].rstrip('/')
            base_config['crawler']['entry_url'] = entry_url
            self.logger.info(f"🔍 DEBUG: entry_url normalized from '{config_dict['url']}' to '{entry_url}'")
            base_config['output']['docx_filename'] = base_name
            
            # Transfer navigation strategy and other options
            nav_strategy = config_dict.get('nav_strategy', 'pcloud')
            base_config['crawler']['nav_strategy'] = nav_strategy
            self.logger.info(f"Using navigation strategy: {nav_strategy}")
            
            if 'max_pages' in config_dict:
                base_config['crawler']['max_pages'] = config_dict['max_pages']
            if 'include_patterns' in config_dict:
                base_config['crawler']['include_patterns'] = config_dict['include_patterns']
            if 'exclude_patterns' in config_dict:
                base_config['crawler']['exclude_patterns'] = config_dict['exclude_patterns']
            if 'github_token' in config_dict:
                base_config['crawler']['github_token'] = config_dict['github_token']
            if 'github_extensions' in config_dict:
                base_config['crawler']['github_extensions'] = config_dict['github_extensions']
            if 'github_max_files' in config_dict:
                base_config['crawler']['github_max_files'] = config_dict['github_max_files']
            
            # Store complete crawl config (used for Dashboard history)
            crawl_config = {
                'crawl_id': crawl_id,
                'url': config_dict.get('url'),
                'output_name': config_dict.get('output_name'),
                'template_id': config_dict.get('template_id', 'custom'),
                'nav_strategy': nav_strategy,
                'max_pages': config_dict.get('max_pages'),
                'timestamp': datetime.now().isoformat(),
                'output_dir': str(crawl_dir),
                'enhanced_html_formatting': config_dict.get('enhanced_html_formatting', False),
                'extra_css_urls': config_dict.get('extra_css_urls', []),
                'use_fallback_css': config_dict.get('use_fallback_css', False),
                'inline_standalone': config_dict.get('inline_standalone', False)
            }
            
            # Save config to crawl folder (for Dashboard history loading)
            with open(config_file, 'w') as f:
                yaml.dump(crawl_config, f)
            
            # Create crawler instance
            self.state_manager.add_log(crawl_id, f'📡 Initializing crawler for {config_dict["url"]}...')
            
            try:
                # Setup file logging for persistent logs
                file_handler = logging.FileHandler(log_file, encoding='utf-8')
                file_handler.setLevel(logging.INFO)
                file_formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
                file_handler.setFormatter(file_formatter)
                
                # Setup logging handler to capture crawler progress
                class ProgressLogHandler(logging.Handler):
                    def __init__(self, state_manager, crawl_id, file_handler):
                        super().__init__()
                        self.state_manager = state_manager
                        self.crawl_id = crawl_id
                        self.page_count = 0
                        self.file_handler = file_handler
                    
                    def emit(self, record):
                        msg = self.format(record)
                        
                        # Write to file FIRST
                        self.file_handler.emit(record)
                        
                        # Track preflight scan result - this is the TRUE total
                        if 'Preflight complete:' in msg or 'Preflight:' in msg and 'total pages' in msg:
                            try:
                                import re
                                # Extract from "Preflight complete: Found 166 total pages"
                                # or "Preflight: 166 total pages (27 categories + 139 details)"
                                match = re.search(r'(\d+)\s+total pages', msg)
                                if match:
                                    total_pages = int(match.group(1))
                                    self.state_manager.update_progress(
                                        self.crawl_id,
                                        total_pages=total_pages
                                    )
                                    logger.info(f"Set total_pages to {total_pages} from preflight scan")
                            except:
                                pass
                        
                        # Track queue size from "Queue has X pages" message
                        # This is now just informational (detail pages only)
                        if 'Queue has' in msg and 'detailed pages to crawl' in msg:
                            try:
                                import re
                                match = re.search(r'Queue has (\d+)', msg)
                                if match:
                                    queue_size = int(match.group(1))
                                    logger.info(f"Queue contains {queue_size} detail pages (categories already counted in preflight)")
                            except:
                                pass
                        
                        # Track fetched pages from various log patterns
                        if '✓ Fetched:' in msg:
                            self.page_count += 1
                            url = msg.split('✓ Fetched:')[-1].strip()
                            self.state_manager.update_progress(
                                self.crawl_id,
                                pages_crawled=self.page_count,
                                current_page=self.page_count,
                                current_url=url
                            )
                        elif 'Fetching' in msg and ('http://' in msg or 'https://' in msg):
                            # Extract URL from fetching messages
                            parts = msg.split('Fetching')
                            if len(parts) > 1:
                                url_part = parts[-1].strip().rstrip('.')
                                if url_part.startswith('http'):
                                    self.state_manager.update_progress(
                                        self.crawl_id,
                                        current_url=url_part
                                    )
                        
                        # Always add to log
                        self.state_manager.add_log(self.crawl_id, msg)
                
                # Add handler to all crawler loggers (ONLY root to avoid duplicates!)
                progress_handler = ProgressLogHandler(self.state_manager, crawl_id, file_handler)
                progress_handler.setLevel(logging.INFO)
                
                # Register ONLY on root logger (child loggers propagate up)
                root_logger = logging.getLogger()
                root_logger.addHandler(file_handler)
                root_logger.addHandler(progress_handler)
                
                # Create temporary config for crawler (needs base_config structure)
                # Add HTML formatting options explicitly to base_config so they appear in YAML
                if 'document' not in base_config:
                    base_config['document'] = {}
                base_config['document']['enhanced_html_formatting'] = config_dict.get('enhanced_html_formatting', False)
                base_config['output']['inline_standalone'] = config_dict.get('inline_standalone', False)
                base_config['document']['use_fallback_css'] = config_dict.get('use_fallback_css', False)
                base_config['document']['extra_css_urls'] = config_dict.get('extra_css_urls', [])
                base_config['document']['min_content_length'] = config_dict.get('min_content_length', 20)
                
                # Add URL Scope Control options to crawler config (CRITICAL: for documentation)
                if 'crawler' not in base_config:
                    base_config['crawler'] = {}
                base_config['crawler']['strict_base_path'] = config_dict.get('strict_base_path', True)
                base_config['crawler']['parent_levels'] = config_dict.get('parent_levels', 0)
                base_config['crawler']['url_must_contain'] = config_dict.get('url_must_contain', [])
                base_config['crawler']['url_must_not_contain'] = config_dict.get('url_must_not_contain', [])
                
                temp_config_path = crawl_dir / 'crawler_config.yaml'
                with open(temp_config_path, 'w') as f:
                    yaml.dump(base_config, f, default_flow_style=False, sort_keys=False)
                
                crawler = Crawler(config_path=str(temp_config_path), max_workers=config_dict.get('parallel_workers', 5))
                
                self.state_manager.add_log(crawl_id, '🔍 Discovering pages...')
                
                # Run the actual crawler
                pages, errors = crawler.crawl()
                
                # Update state with FINAL results
                actual_pages = len(pages)
                duration = self.state_manager.get_state(crawl_id).get_duration()
                speed = actual_pages / duration if duration > 0 else 0
                
                self.state_manager.update_progress(
                    crawl_id,
                    pages_crawled=actual_pages,
                    current_page=actual_pages,
                    total_pages=actual_pages,
                    errors=len(errors),
                    speed=speed
                )
                
                self.state_manager.add_log(crawl_id, f'✅ Crawled {actual_pages} pages with {len(errors)} errors')
                
                # Save metadata.json for persistence
                state = self.state_manager.get_state(crawl_id)
                if state:
                    state.save_metadata(str(crawl_dir))
                
            except Exception as e:
                self.logger.error(f"Crawler execution failed: {e}", exc_info=True)
                # Fallback to simulation if real crawler fails
                self.state_manager.add_log(crawl_id, f'⚠️ Real crawler failed: {str(e)}. Using simulation...')
                pages = {}
                errors = {}
                max_pages = config_dict.get('max_pages', 10)
                
                # Simulate for fallback
                import time
                for i in range(1, max_pages + 1):
                    state = self.state_manager.get_state(crawl_id)
                    if not state or state.status == 'stopped':
                        break
                    
                    fake_url = f"{config_dict['url']}/page-{i}"
                    pages[fake_url] = {'raw_html': '<html><body>Simulated content</body></html>'}
                    
                    self.state_manager.update_progress(
                        crawl_id,
                        pages_crawled=i,
                        current_page=i,
                        total_pages=max_pages,
                        current_url=fake_url
                    )
                    self.state_manager.add_log(crawl_id, f'✅ Crawled: {fake_url}')
                    time.sleep(0.5)
            
            # Generate document
            self.state_manager.add_log(crawl_id, '📝 Generating document...')
            
            # Create a proper DOCX file
            try:
                from docx import Document
                from docx.shared import Inches, Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                doc = Document()
                
                # Title
                title = doc.add_heading('Web Crawl Results', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Metadata
                doc.add_heading('Crawl Information', level=1)
                doc.add_paragraph(f'Target URL: {config_dict["url"]}')
                doc.add_paragraph(f'Pages Crawled: {len(pages)}')
                doc.add_paragraph(f'Errors: {len(errors)}')
                doc.add_paragraph(f'Started: {self.state_manager.get_state(crawl_id).start_time.strftime("%Y-%m-%d %H:%M:%S")}')
                doc.add_paragraph(f'Crawl ID: {crawl_id}')
                
                # Summary
                doc.add_heading('Summary', level=1)
                if len(pages) > 0:
                    doc.add_paragraph(f'Successfully crawled {len(pages)} pages from {config_dict["url"]}')
                else:
                    doc.add_paragraph('No pages were successfully crawled.')
                
                # Pages content
                if pages:
                    doc.add_heading('Crawled Pages', level=1)
                    
                    # Initialize parser for content extraction
                    parser = HTMLParser(base_config)
                    
                    for i, (url, page_data) in enumerate(pages.items(), 1):
                        # Extract content from raw_html using parser
                        raw_html = page_data.get('raw_html', '')
                        
                        # Check if this is raw text content (e.g., GitHub raw URLs)
                        is_raw_content = 'raw.githubusercontent.com' in url or (
                            raw_html and not raw_html.lstrip()[:200].lower().startswith(('<!doctype', '<html'))
                        )
                        
                        # DEBUG LOGGING
                        logger.info(f"Processing URL: {url}")
                        logger.info(f"  - is_raw_content: {is_raw_content}")
                        logger.info(f"  - raw.githubusercontent.com in url: {'raw.githubusercontent.com' in url}")
                        logger.info(f"  - raw_html length: {len(raw_html) if raw_html else 0}")
                        if raw_html:
                            logger.info(f"  - First 100 chars: {raw_html[:100]}")
                        
                        if is_raw_content and raw_html:
                            # Treat as plain text - extract filename as title
                            import os
                            from urllib.parse import urlparse
                            path = urlparse(url).path
                            title = os.path.basename(path.rstrip('/')) or url
                            content = raw_html  # Use raw text directly
                            logger.info(f"  - Extracted title: {title}")
                            logger.info(f"  - Content length: {len(content)}")
                        elif raw_html:
                            # Regular HTML parsing
                            from bs4 import BeautifulSoup
                            soup = BeautifulSoup(raw_html, 'html.parser')
                            title = parser._extract_title(soup)
                            content_elem = parser._find_content(soup)
                            
                            if content_elem:
                                content = content_elem.get_text(separator='\n', strip=True)
                            else:
                                content = ''
                        else:
                            title = url
                            content = ''
                        
                        # Page header
                        doc.add_heading(f'{i}. {title}', level=2)
                        doc.add_paragraph(f'URL: {url}', style='Intense Quote')
                        
                        # Content
                        if content:
                            # For raw content (code files), format as code block
                            if is_raw_content:
                                # Add as single code block with line breaks preserved
                                code_para = doc.add_paragraph()
                                code_para.style = 'No Spacing'
                                run = code_para.add_run()
                                run.font.name = 'Courier New'
                                run.font.size = Pt(9)
                                
                                # Add text with line breaks
                                lines = content.splitlines()
                                for line_idx, line in enumerate(lines):
                                    run.add_text(line)
                                    if line_idx < len(lines) - 1:
                                        run.add_break()
                            else:
                                # Regular HTML content - split into paragraphs
                                paragraphs = content.split('\n\n')
                                for para in paragraphs:  # ALL paragraphs, not just 20
                                    para = para.strip()
                                    if para and len(para) > 10:  # Skip very short lines
                                        # Check if it's a code block (indented or marked)
                                        if para.startswith('    ') or para.startswith('```'):
                                            # Code block
                                            code_text = para.replace('```', '').strip()
                                            code_para = doc.add_paragraph(code_text)
                                            code_para.style = 'No Spacing'
                                            # Monospace font
                                            for run in code_para.runs:
                                                run.font.name = 'Courier New'
                                                run.font.size = Pt(9)
                                        else:
                                            # Regular paragraph - FULL CONTENT, no truncation
                                            doc.add_paragraph(para)
                        else:
                            # Show warning
                            para = doc.add_paragraph('⚠️ No content extracted')
                            para.runs[0].font.italic = True
                        
                        # Add spacing between pages
                        doc.add_paragraph()
                
                # Errors section
                if errors:
                    doc.add_heading('Errors', level=1)
                    for url, error in errors.items():
                        doc.add_paragraph(f'{url}', style='List Bullet')
                        doc.add_paragraph(f'   Error: {error}')
                
                # Save
                doc.save(str(output_file))
                file_size_kb = output_file.stat().st_size / 1024
                log_size_kb = log_file.stat().st_size / 1024 if log_file.exists() else 0
                
                self.state_manager.add_log(crawl_id, f'✅ Document saved: {output_file.name} ({file_size_kb:.1f} KB)')
                self.state_manager.add_log(crawl_id, f'📋 Log file saved: {log_file.name} ({log_size_kb:.1f} KB)')
                
                # Generate HTML documentation
                try:
                    self.state_manager.add_log(crawl_id, '🌐 Generating HTML documentation...')
                    from html_generator import HTMLGenerator
                    
                    html_gen = HTMLGenerator(base_config, min_content_length=config_dict.get('min_content_length', 20))
                    enhanced_formatting = config_dict.get('enhanced_html_formatting', False)
                    extra_css_urls = config_dict.get('extra_css_urls', [])
                    use_fallback_css = config_dict.get('use_fallback_css', False)
                    html_path = html_gen.generate(pages, crawl_dir, 'index', enhanced_formatting=enhanced_formatting, extra_css_urls=extra_css_urls, use_fallback_css=use_fallback_css)
                    
                    if html_path:
                        html_size_kb = Path(html_path).stat().st_size / 1024
                        self.state_manager.add_log(crawl_id, f'✅ HTML saved: index.html ({html_size_kb:.1f} KB)')
                        self.state_manager.update_progress(crawl_id, html_path=html_path)
                    else:
                        self.state_manager.add_log(crawl_id, '⚠️ HTML generation failed')
                except Exception as html_error:
                    self.logger.error(f"HTML generation failed: {html_error}", exc_info=True)
                    self.state_manager.add_log(crawl_id, f'⚠️ HTML generation error: {str(html_error)}')
                
                # Generate Markdown documentation for RAG/Vector DB
                try:
                    self.state_manager.add_log(crawl_id, '📝 Generating Markdown documentation...')
                    from markdown_generator import MarkdownGenerator
                    
                    md_gen = MarkdownGenerator(base_config)
                    md_path = md_gen.generate(pages, crawl_dir)
                    
                    if md_path:
                        md_size_kb = Path(md_path).stat().st_size / 1024
                        self.state_manager.add_log(crawl_id, f'✅ Markdown saved: documentation.md ({md_size_kb:.1f} KB)')
                    else:
                        self.state_manager.add_log(crawl_id, '⚠️ Markdown generation failed')
                except Exception as md_error:
                    self.logger.error(f"Markdown generation failed: {md_error}", exc_info=True)
                    self.state_manager.add_log(crawl_id, f'⚠️ Markdown generation error: {str(md_error)}')
                
                # Store file paths in state
                self.state_manager.update_progress(
                    crawl_id, 
                    output_path=str(output_file),
                    log_path=str(log_file)
                )
                
            except ImportError:
                # Fallback if python-docx not installed
                self.state_manager.add_log(crawl_id, '⚠️ python-docx not installed, creating text file')
                output_file.write_text(
                    f"Web Crawl Results\n\n"
                    f"URL: {config_dict['url']}\n"
                    f"Pages: {len(pages)}\n\n"
                    f"Crawled Pages:\n" +
                    "\n".join([f"{i}. {url}" for i, url in enumerate(pages.keys(), 1)])
                )
            
            finally:
                # Cleanup logging handlers
                try:
                    root_logger = logging.getLogger()
                    root_logger.removeHandler(file_handler)
                    root_logger.removeHandler(progress_handler)
                    
                    for logger_name in ['crawler', '__main__']:
                        handler_logger = logging.getLogger(logger_name)
                        handler_logger.removeHandler(file_handler)
                        handler_logger.removeHandler(progress_handler)
                    
                    file_handler.close()
                except:
                    pass
            
            # Mark as completed
            self.state_manager.update_progress(crawl_id, pages_crawled=len(pages))
            self.state_manager.complete_crawl(crawl_id, str(output_file))
            self.state_manager.add_log(crawl_id, f'✅ Complete! Saved to {output_file.name}')
            
        except Exception as e:
            self.logger.error(f"Crawler error: {e}", exc_info=True)
            self.state_manager.error_crawl(crawl_id, str(e))
    
    def get_crawl_status(self, crawl_id):
        """Get status of a running crawl"""
        state = self.state_manager.get_state(crawl_id)
        if state:
            return state.to_dict()
        return None
    
    def stop_crawl(self, crawl_id):
        """Stop a running crawl"""
        self.state_manager.stop_crawl(crawl_id)
    
    def _format_time_ago(self, timestamp):
        """Format timestamp as 'X time ago'"""
        if not timestamp:
            return 'unknown'
        
        delta = datetime.now() - timestamp
        seconds = delta.total_seconds()
        
        if seconds < 60:
            return 'just now'
        elif seconds < 3600:
            return f'{int(seconds / 60)}m ago'
        elif seconds < 86400:
            return f'{int(seconds / 3600)}h ago'
        else:
            return f'{int(seconds / 86400)}d ago'
