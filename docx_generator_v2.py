import logging
from typing import Dict, List, Any, Optional
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE

from image_handler import ImageHandler
from syntax_highlighter import add_highlighted_code_to_paragraph
from template_manager import TemplateManager

logger = logging.getLogger(__name__)


class DOCXGeneratorV2:
    def __init__(self, config: Dict[str, Any]):
        self.config = config.get('document', {})
        self.doc = Document()
        self.image_handler = ImageHandler()
        self.base_url = config.get('crawler', {}).get('entry_url', 'https://docs.pcloud.com/')
        
        template_name = self.config.get('template', 'standard')
        self.template = TemplateManager.get_template(template_name)
        self.template_config = self.template.get_config()
        
        logger.info(f"Using template: {self.template.name}")
        
        self._setup_styles()
    
    def _setup_styles(self):
        """Setup document styles"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
    
    def _add_hyperlink(self, paragraph, text: str, url: str):
        """Add clickable hyperlink to paragraph"""
        try:
            part = paragraph.part
            r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
            
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            
            rStyle = OxmlElement('w:rStyle')
            rStyle.set(qn('w:val'), 'Hyperlink')
            rPr.append(rStyle)
            new_run.append(rPr)
            
            text_run = OxmlElement('w:t')
            text_run.text = text
            new_run.append(text_run)
            
            hyperlink.append(new_run)
            paragraph._p.append(hyperlink)
        except Exception as e:
            logger.error(f"Error adding hyperlink: {e}")
            paragraph.add_run(text)
    
    def _add_image(self, img_url: str, alt_text: str = '', width: float = 5.0) -> bool:
        """
        Download and add image to document.
        
        Args:
            img_url: Image URL (relative or absolute)
            alt_text: Alternative text for image
            width: Width in inches
            
        Returns:
            True if image was added, False if failed
        """
        try:
            img_path = self.image_handler.download_image(img_url, self.base_url)
            if img_path and img_path.exists():
                self.doc.add_picture(str(img_path), width=Inches(width))
                if alt_text:
                    caption = self.doc.add_paragraph(f"Figure: {alt_text}", style='Caption')
                    caption.runs[0].font.size = Pt(9)
                return True
        except Exception as e:
            logger.warning(f"Failed to add image {img_url}: {e}")
        
        return False
    
    def _add_code_block(self, code: str, language: Optional[str] = None):
        """
        Add syntax-highlighted code block.
        
        Args:
            code: Code content
            language: Programming language (python, json, etc)
        """
        if language:
            lang_para = self.doc.add_paragraph(f"Code ({language})")
            lang_para.runs[0].font.size = Pt(9)
            lang_para.runs[0].font.italic = True
            lang_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        code_para = self.doc.add_paragraph()
        code_para.style = 'Intense Quote'
        code_para.paragraph_format.left_indent = Inches(0.25)
        
        add_highlighted_code_to_paragraph(code_para, code, language)
    
    def _add_title_page(self):
        """Add professional title page"""
        title = self.doc.add_heading('pCloud API Documentation', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        
        subtitle = self.doc.add_heading('Compiled Reference', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        source = self.doc.add_paragraph('Source: docs.pcloud.com')
        source.alignment = WD_ALIGN_PARAGRAPH.CENTER
        source.runs[0].font.size = Pt(10)
        
        date_para = self.doc.add_paragraph(f'Generated: {datetime.now().strftime("%d.%m.%Y")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.runs[0].font.size = Pt(10)
        
        self.doc.add_page_break()
    
    def _add_toc_placeholder(self):
        """Add real TOC field"""
        toc_heading = self.doc.add_heading('Table of Contents', level=1)
        
        toc_para = self.doc.add_paragraph()
        run = toc_para.add_run()
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        note = self.doc.add_paragraph('(Right-click and select "Update Field" to refresh)')
        note.runs[0].font.italic = True
        note.runs[0].font.size = Pt(9)
        note.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        self.doc.add_page_break()
    
    def _add_category_overview(self, category: str, methods: List[Dict[str, Any]]):
        """Add overview section for a category"""
        overview = self.doc.add_heading(category, level=1)
        
        description = f"This section contains the {category} API methods."
        self.doc.add_paragraph(description)
        
        overview_heading = self.doc.add_heading('Methods Overview', level=2)
        overview_heading.runs[0].font.size = Pt(12)
        
        for method in methods:
            name = method.get('name', 'Unknown')
            api_url = method.get('api_url', '')
            
            bullet = self.doc.add_paragraph(f'{name} – ', style='List Bullet')
            bullet.runs[0].font.size = Pt(10)
            if api_url:
                self._add_hyperlink(bullet, api_url, api_url)
        
        self.doc.add_paragraph()
    
    def _add_method_section(self, method: Dict[str, Any]):
        """Add a method documentation section"""
        name = method.get('name', 'Unknown')
        
        method_heading = self.doc.add_heading(name, level=2)
        
        auth = method.get('auth', 'unknown')
        api_url = method.get('api_url', '')
        doc_url = method.get('url', '')
        
        info_table = self.doc.add_table(rows=3, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        info_cells = info_table.rows[0].cells
        info_cells[0].text = 'Auth'
        info_cells[1].text = auth
        
        info_cells = info_table.rows[1].cells
        info_cells[0].text = 'API URL'
        info_cells[1].text = ''
        url_para = info_cells[1].paragraphs[0]
        self._add_hyperlink(url_para, api_url, api_url)
        
        info_cells = info_table.rows[2].cells
        info_cells[0].text = 'Doc Page'
        info_cells[1].text = ''
        doc_para = info_cells[1].paragraphs[0]
        self._add_hyperlink(doc_para, doc_url, doc_url)
        
        self.doc.add_paragraph()
        
        description = method.get('description', '')
        if description:
            desc_heading = self.doc.add_heading('Description', level=3)
            self.doc.add_paragraph(description)
        
        notes = method.get('notes', [])
        if notes:
            notes_heading = self.doc.add_heading('Notes', level=3)
            for note in notes:
                self.doc.add_paragraph(note, style='List Bullet')
        
        parameters = method.get('parameters', {})
        if parameters.get('required') or parameters.get('optional'):
            self._add_parameters_section(parameters)
        
        output = method.get('output', [])
        if output:
            self._add_output_section(output)
        
        example = method.get('example', '')
        if example:
            example_heading = self.doc.add_heading('Example', level=3)
            self._add_code_block(example, language='json')
        
        errors = method.get('errors', [])
        if errors:
            self._add_errors_section(errors)
        
        self.doc.add_paragraph()
        self.doc.add_paragraph()
    
    def _add_parameters_section(self, parameters: Dict[str, List[Dict]]):
        """Add parameters section with tables"""
        param_heading = self.doc.add_heading('Parameters', level=3)
        
        required = parameters.get('required', [])
        optional = parameters.get('optional', [])
        
        if required:
            self.doc.add_heading('Required', level=4)
            self._add_parameter_table(required)
        
        if optional:
            self.doc.add_heading('Optional', level=4)
            self._add_parameter_table(optional)
    
    def _add_parameter_table(self, params: List[Dict]):
        """Add parameter table"""
        if not params:
            return
        
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Name'
        header_cells[1].text = 'Type'
        header_cells[2].text = 'Description'
        
        for param in params:
            row_cells = table.add_row().cells
            row_cells[0].text = param.get('name', '')
            row_cells[1].text = param.get('type', '')
            row_cells[2].text = param.get('description', '')
        
        self.doc.add_paragraph()
    
    def _add_output_section(self, output: List[Dict]):
        """Add output/response section"""
        if not output:
            return
        
        output_heading = self.doc.add_heading('Output', level=3)
        
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Field'
        header_cells[1].text = 'Type'
        header_cells[2].text = 'Description'
        
        for field in output:
            row_cells = table.add_row().cells
            row_cells[0].text = field.get('parameter', field.get('field', ''))
            row_cells[1].text = field.get('type', '')
            row_cells[2].text = field.get('description', '')
        
        self.doc.add_paragraph()
    
    def _add_errors_section(self, errors: List[Dict]):
        """Add errors section"""
        if not errors:
            return
        
        errors_heading = self.doc.add_heading('Errors', level=3)
        
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Code'
        header_cells[1].text = 'Description'
        
        for error in errors:
            row_cells = table.add_row().cells
            code = error.get('code', '')
            row_cells[0].text = code
            row_cells[1].text = error.get('description', '')
        
        self.doc.add_paragraph()
    
    def generate(self, categories: Dict[str, List[Dict[str, Any]]]):
        """Generate DOCX from categorized methods"""
        logger.info("Generating categorized DOCX document...")
        
        self._add_title_page()
        self._add_toc_placeholder()
        
        total_methods = 0
        
        for category, methods in categories.items():
            logger.info(f"Adding category: {category} ({len(methods)} methods)")
            
            self._add_category_overview(category, methods)
            
            for method in methods:
                self._add_method_section(method)
                total_methods += 1
        
        logger.info(f"Added {total_methods} method pages")
        
        return total_methods
    
    def save(self, output_path: str):
        """Save document"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        self.doc.save(str(output_path))
        logger.info(f"Document saved to {output_path}")
